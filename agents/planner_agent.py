# agents/planner_agent.py
"""
PLANNER AGENT
=============
RESPONSIBILITY: Orchestrates the creation of comprehensive history lesson packages.

PURPOSE:
- Plans multi-lesson units on historical topics
- Coordinates research, content generation, and file creation
- Manages Worker, Fact-Checker, and PPT agents
- Produces Teacher Guides (DOCX) and PowerPoint presentations

FLOW:
1. Receives validated history request from request_reviewer
2. Creates lesson plan structure
3. Coordinates research via Worker agent (Britannica by default, Wikipedia fallback)
4. Generates detailed Teacher's Guide content
5. Creates PowerPoint slide structures
6. Dispatches to PPT agent for file generation
7. Returns file paths to user

GUARDRAILS:
- Assumes all requests are pre-validated as history-related
- Focuses strictly on historical education content
- Uses evidence-based content generation
"""

import asyncio
import json
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List

try:
    import docx
except ImportError:
    docx = None

from queues.message_bus import task_queue, result_queue, ppt_queue
from utils.llm import generate
from agents.worker_agent import run_worker_step as worker_agent
from agents.fact_checker_agent import fact_checker_agent
from agents.quizzer_agent import generate_quiz, format_quiz_for_docx


OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Config
SLIDE_TARGET = 30
MAX_WIKI_TOPICS = 5
DEFAULT_RESEARCH_TOOL = "britannica"
FACT_CHECK_ENABLED = True  # Enable LLM-based fact checking


def _slugify_title(title: str) -> str:
    """Return a filesystem-safe slug derived from the desired PPT title."""
    safe_text = (title or "presentation").encode("ascii", errors="ignore").decode()
    safe_text = re.sub(r"[^A-Za-z0-9]+", "-", safe_text).strip("-") or "presentation"
    return safe_text[:60].lower()


def build_ppt_filename(title: str) -> str:
    """Create a unique PPT filename."""
    base_slug = _slugify_title(title)
    candidate = base_slug
    counter = 1
    while (OUTPUT_DIR / f"{candidate}.pptx").exists():
        candidate = f"{base_slug}-{counter}"
        counter += 1
    return f"{candidate}.pptx"


def _safe_json_loads(raw: str) -> Dict[str, Any]:
    if not raw:
        return {}
    raw = raw.strip()
    match = re.search(r"```json\s*(.*?)\s*```", raw, flags=re.DOTALL)
    if match:
        raw = match.group(1)
    
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass
    return {}

# -------------------------------------------------------------------------
#  DOCX HELPER
# -------------------------------------------------------------------------
def create_docx(filename: str, title: str, summary_text: str):
    """Creates a Word document with continuous flowing text."""
    if not docx:
        print("[Planner] python-docx not installed, skipping DOCX generation.")
        return False
        
    try:
        doc = docx.Document()
        doc.add_heading(title, 0)
        
        # Add the summary as continuous text, preserving paragraph breaks
        paragraphs = summary_text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Remove any markdown formatting for cleaner continuous text
                para_text = para_text.replace('**', '').replace('##', '').replace('#', '')
                doc.add_paragraph(para_text)
        
        out_path = OUTPUT_DIR / filename
        doc.save(out_path)
        return True
    except Exception as e:
        print(f"[Planner] Error writing DOCX: {e}")
        return False


def create_sources_document(filename: str, unit_title: str, sources: List[Dict[str, str]]):
    """Creates a sources document listing all research materials used."""
    if not docx:
        print("[Planner] python-docx not installed, skipping sources document.")
        return False
        
    try:
        doc = docx.Document()
        doc.add_heading(f"Sources for: {unit_title}", 0)
        
        doc.add_paragraph(
            "This document lists all sources consulted during the creation of this lesson unit. "
            "All content was fact-checked against these sources using Natural Language Inference verification."
        )
        
        doc.add_heading("Research Sources", level=1)
        
        # Group sources by lesson
        current_lesson = None
        for source in sources:
            lesson = source.get("lesson", "Unknown Lesson")
            if lesson != current_lesson:
                doc.add_heading(lesson, level=2)
                current_lesson = lesson
            
            source_type = source.get("type", "Article")
            title = source.get("title", "Untitled")
            url = source.get("url", "")
            topic = source.get("topic", "")
            
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{source_type}: ").bold = True
            p.add_run(f"{title}")
            if topic:
                p.add_run(f" (Topic: {topic})")
            if url:
                p.add_run(f"\n   {url}")
        
        out_path = OUTPUT_DIR / filename
        doc.save(out_path)
        print(f"[Planner] Sources document saved: {out_path}")
        return True
    except Exception as e:
        print(f"[Planner] Error creating sources document: {e}")
        return False

# -------------------------------------------------------------------------
#  PROMPTS
# -------------------------------------------------------------------------

def _plan_lessons_prompt(topic: str, num_lessons: int) -> str:
    return f"""
You are an expert history teacher creating educational content.

**CRITICAL REQUIREMENT: This is a HISTORY education system.**
You must create a comprehensive {num_lessons}-lesson unit plan on: "{topic}".

Return STRICT JSON matching this schema:
{{
  "unit_title": "Title of the whole unit",
  "lessons": [
    {{
      "lesson_number": 1,
      "title": "Concrete Topic Name",
      "topics_to_research_on_britannica": ["Topic 1", "Topic 2"]
    }}
  ]
}}

MANDATORY Requirements:
- EXACTLY {num_lessons} lessons
- HISTORY-FOCUSED: All lessons must cover historical events, periods, figures, or movements
- Lesson titles must be CONCRETE and FACTUAL (e.g. "The Storming of the Bastille", not "A Dream of Liberty")
- Number the lessons sequentially
- Maintain academic rigor appropriate for history education

**CRITICAL: Research Topics MUST be SPECIFIC and ACCURATE:**
- Include specific dates, event names, or proper nouns (e.g., "Korean War 1950-1953", NOT "Korean Peninsula")
- Use exact names of treaties, battles, or events (e.g., "Battle of Inchon", NOT "Maoke Agreement")
- Include the full context (e.g., "Cold War origins 1945-1950", NOT just "Cold War")
- Verify that topics are real historical subjects that will have Britannica articles
- For wars: include dates and key battles/events (e.g., "World War II D-Day invasion", "Korean War causes")
- For people: use full names and roles (e.g., "Winston Churchill World War II", NOT just "Churchill")
- AVOID vague geographic or period terms alone (e.g., NOT "Korean Peninsula", NOT "Industrial Age")

Examples of GOOD research topics:
- "Korean War 1950-1953"
- "38th parallel Korea division"
- "Inchon Landing September 1950"
- "French Revolution storming of Bastille"
- "Treaty of Versailles 1919"

Examples of BAD research topics:
- "Korean Peninsula" (too vague, will return ancient history)
- "Cold War" (too broad, won't return specific events)
- "Revolution period" (not specific enough)
""".strip()


def _teacher_summary_prompt(lesson_title: str, evidence: str, previous_content: str = "") -> str:
    previous_context = ""
    if previous_content:
        previous_context = f"""
CONTEXT FROM PREVIOUS LESSONS:
{previous_content}

CRITICAL: The above content was already covered in previous lessons. 
DO NOT repeat these topics, facts, or concepts. Build upon them and introduce NEW information.
Reference previous lessons briefly if needed for continuity, but focus on ORIGINAL content.

"""
    
    return f"""
**SYSTEM GUARDRAIL: HISTORY EDUCATION ONLY**
You are an expert HISTORY teacher writing a comprehensive "Teacher's Guide" for: "{lesson_title}".

This guide will be saved as a Word document and accompanies a 30-slide PowerPoint presentation.
The guide must contain ALL the detailed knowledge the teacher needs, written as CONTINUOUS FLOWING TEXT.

{previous_context}IMPORTANT STRUCTURAL REQUIREMENTS:
1. Write in continuous paragraphs (NOT bullet points or lists)
2. Organize content into EXACTLY 30 logical sections (one per slide)
3. Each section should be 3-5 sentences that expand on one key concept
4. Start each section with a clear topic sentence that could serve as a slide title
5. Use academic, clear language - NO poetic or "dreamy" phrasing
6. Ensure all content is NEW and does not repeat what was covered before

**HISTORY GUARDRAILS:**
- Focus EXCLUSIVELY on historical facts, events, and analysis
- Base all content on the provided evidence from historical sources
- Maintain chronological accuracy and historical context
- Use proper historical terminology and period-appropriate references
- Cite specific dates, figures, and locations where relevant

Format:
[Section 1 topic]
[3-5 sentences of detailed explanation]

[Section 2 topic]
[3-5 sentences of detailed explanation]

... continue for 30 sections total

EVIDENCE:
{evidence}
""".strip()


def _extract_notes_from_summary(teacher_summary: str, slides: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Extracts notes from the teacher summary by matching sections to slide titles.
    Each slide gets the corresponding section from the summary as notes.
    """
    # Split summary into sections (paragraphs separated by double newlines)
    sections = [s.strip() for s in teacher_summary.split('\n\n') if s.strip()]
    
    # Match slides to sections (assuming they're in order)
    for i, slide in enumerate(slides):
        if i < len(sections):
            slide['notes'] = sections[i]
        else:
            slide['notes'] = ''
    
    return slides


def _slide_generation_prompt(lesson_title: str, teacher_summary: str, target_slide_count: int, previous_slide_titles: List[str] = None) -> str:
    previous_context = ""
    if previous_slide_titles:
        previous_context = f"""
SLIDE TITLES FROM PREVIOUS LESSONS:
{chr(10).join(f'- {title}' for title in previous_slide_titles)}

CRITICAL: Avoid creating slides with similar titles or covering the same topics.
Your slides must introduce NEW concepts and information.

"""
    
    return f"""
You are an instructional designer. Create a PowerPoint structure for: "{lesson_title}".
Target length: EXACTLY {target_slide_count} slides.

The Teacher's Guide is organized into {target_slide_count} sections. Create ONE SLIDE PER SECTION.

{previous_context}Return STRICT JSON:
{{
  "slides": [
    {{
      "title": "Slide Title",
      "bullets": ["Keyword 1", "Keyword 2", "Short phrase"],
      "is_question": false
    }},
    {{
      "title": "🤔 Think About It",
      "bullets": ["Why would X be significant?", "What might have happened if Y?"],
      "is_question": true
    }}
  ]
}}

MANDATORY RULES:
1. Extract the topic sentence from each section as the slide title
2. Convert key facts from each section into 2-4 bullets
3. Bullets MUST be keywords, dates, names, or short phrases (<= 6 words)
4. NO complete sentences in bullets (EXCEPT for question slides)
5. Maintain the EXACT order of sections from the guide
6. Slide titles must be concrete and factual (e.g., "Economic Crisis 1789")
7. You MUST generate EXACTLY {target_slide_count} slides - no more, no less
8. Ensure content is unique and does not duplicate previous lessons

**QUESTION SLIDES - CRITICAL REQUIREMENT:**
- Include EXACTLY 2 "question slides" in the lesson
- Question slides MUST appear at positions 10 and 20 (after the 10th and 20th content slides)
- For question slides:
  * Set "is_question": true
  * Use title: "🤔 Think About It" or "💭 Critical Thinking Question"
  * Each question slide should have EXACTLY ONE thought-provoking question in the bullets
  * Questions should make students think LOGICALLY and analytically
  * Questions do NOT need to be directly answerable from the content presented so far
  * Examples: "Why would coal and steel be the two important resources to focus on in the ECSC after WW2?"
              "Why do you think the Treaty of Versailles led to further conflict?"
              "Why do you think did the nations of Europe oppose the French Revolution?"
  * Questions should encourage deeper historical thinking about causes, effects, and significance

**HISTORY EDUCATION REQUIREMENTS:**
- All slide content must be historically accurate and factual
- Include specific dates, names, and locations in bullets
- Focus on cause-effect relationships and historical significance
- Maintain educational value appropriate for history teaching
- Notes should contain the complete detailed text from the corresponding Teacher's Guide section
- Question slides should relate to the surrounding content but challenge students to think beyond the facts

TEACHER'S GUIDE:
{teacher_summary}
""".strip()


# -------------------------------------------------------------------------
#  WORKFLOW STEPS
# -------------------------------------------------------------------------

async def _gather_evidence(
    topics: List[str],
    evidence_cache: Dict[str, str],
    research_tool: str = DEFAULT_RESEARCH_TOOL,
) -> str:
    """
    Gathers evidence for topics, using cache to avoid duplicate research.
    Updates the evidence_cache with newly researched topics.
    """
    evidence_parts = []
    unique_topics = []
    seen = set()
    for t in topics:
        clean = t.strip()
        if clean.lower() not in seen:
            seen.add(clean.lower())
            unique_topics.append(clean)
    
    unique_topics = unique_topics[:MAX_WIKI_TOPICS]

    for topic in unique_topics:
        # Check cache first
        cache_key = topic.lower()
        if cache_key in evidence_cache:
            print(f"[Planner] Using cached evidence for: {topic}")
            evidence_parts.append(f"--- article: {topic} (cached) ---\n{evidence_cache[cache_key]}\n")
            continue
            
        # Research if not cached
        print(f"[Planner] Researching: {topic}")
        step_cmd = f"TOOL:{research_tool}:{topic}"
        result = await worker_agent(step_cmd)
        
        # Store in cache
        evidence_cache[cache_key] = result
        evidence_parts.append(f"--- article: {topic} ---\n{result}\n")
    
    return "\n".join(evidence_parts)


async def _process_lesson(lesson_info: Dict[str, Any], unit_title: str, shared_context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Handles end-to-end generation for a single lesson.
    Uses shared_context to avoid repetition across lessons.
    Returns a dict with lesson details and file paths.
    """
    l_num = lesson_info.get("lesson_number", "?")
    l_title = lesson_info.get("title", "Untitled")
    topics = lesson_info.get("topics_to_research_on_britannica")
    if topics is None:
        topics = lesson_info.get("topics_to_research_on_wikipedia", [])
    
    # Concrete naming: "Lesson 1 - Topic Name"
    full_lesson_name = f"Lesson {l_num} - {l_title}"
    print(f"\n[Planner] Processing {full_lesson_name}...")

    # 1. Research (using shared evidence cache)
    evidence = await _gather_evidence(topics, shared_context["evidence_cache"])
    if not evidence.strip():
        evidence = "No specific evidence found."
    
    # Temporarily store source candidates (will validate after fact checking)
    temp_sources = []
    article_sections = re.split(r'--- article: (.+?) ---', evidence)
    sources_seen = set()  # Track to avoid duplicates
    
    for i in range(1, len(article_sections), 2):
        if i + 1 < len(article_sections):
            topic_name = article_sections[i].strip()
            article_content = article_sections[i + 1]
            
            # Extract title and URL from this specific article
            title_match = re.search(r'\*\*(?:Encyclopaedia Britannica|Wikipedia) Article Used:\*\*\s*([^\n]+)', article_content)
            url_match = re.search(r'🔗\s*(https?://[^\s]+)', article_content)
            
            source_title = title_match.group(1).strip() if title_match else topic_name
            source_url = url_match.group(1) if url_match else ""
            
            # Create unique identifier to prevent duplicates
            source_key = f"{source_title}|{source_url}"
            
            if source_key not in sources_seen:
                sources_seen.add(source_key)
                temp_sources.append({
                    "lesson": full_lesson_name,
                    "topic": topic_name,
                    "type": "Encyclopaedia Britannica" if "britannica" in article_content.lower() else "Wikipedia",
                    "title": source_title,
                    "url": source_url
                })

    # 2. Write Teacher Summary (The Real Knowledge) - with context from previous lessons
    print(f"[Planner] Writing Teacher Guide for {full_lesson_name}...")
    previous_summaries = "\n\n---PREVIOUS LESSON---\n\n".join(shared_context["lesson_summaries"])
    summary = await generate(
        _teacher_summary_prompt(full_lesson_name, evidence, previous_summaries), 
        max_tokens=2000
    )
    
    # 2b. FACT CHECK the generated summary
    irrelevant_sources = []
    if FACT_CHECK_ENABLED:
        print(f"[Planner] Fact-checking content for {full_lesson_name}...")
        fact_check_result = await fact_checker_agent(summary, evidence)
        print(fact_check_result)
        
        # Parse fact check result to identify irrelevant sources
        if "Warnings:" in fact_check_result:
            warnings_section = fact_check_result.split("Warnings:")[1] if "Warnings:" in fact_check_result else ""
            # Look for mentions of specific article titles or topics in warnings
            for source in temp_sources:
                source_identifier = source["title"].split("|")[0].strip()  # Get first part of title
                # Check if this source is mentioned as irrelevant in warnings
                if source_identifier.lower() in warnings_section.lower() or \
                   "unrelated" in warnings_section.lower() or \
                   "not relevant" in warnings_section.lower():
                    # Check if the source title appears in the warning
                    for word in source_identifier.split():
                        if len(word) > 4 and word.lower() in warnings_section.lower():
                            irrelevant_sources.append(source["title"])
                            print(f"[Planner] 🚫 Excluding irrelevant source: {source['title']}")
                            break
        
        # Revision loop - keep trying until we get GO or hit max attempts
        max_revision_attempts = 4
        revision_attempt = 0
        
        while "GO/NO-GO: GO" not in fact_check_result and revision_attempt < max_revision_attempts:
            # Check if there are actual warnings (not just "None")
            has_warnings = "Warnings:" in fact_check_result and "None" not in fact_check_result.split("Warnings:")[1][:20]
            is_no_go = "GO/NO-GO: NO-GO" in fact_check_result
            
            if not (is_no_go or has_warnings):
                break  # No issues, exit loop
            
            revision_attempt += 1
            print(f"[Planner] ⚠️ Content needs revision (attempt {revision_attempt}/{max_revision_attempts})...")
            
            # Create revision prompt with fact checker feedback
            revision_prompt = f"""
You are an expert history teacher revising educational content based on fact-checker feedback.

**ORIGINAL LESSON:** {full_lesson_name}

**FACT CHECKER FEEDBACK:**
{fact_check_result}

**AVAILABLE EVIDENCE (USE ONLY THIS):**
{evidence}

**YOUR TASK:**
Rewrite the teacher's guide for this lesson, addressing ALL warnings from the fact checker.

**CRITICAL RULES:**
1. Use ONLY information from the provided evidence
2. Remove or correct any unsupported claims mentioned in the warnings
3. Do NOT include content about unrelated topics (e.g., ancient history when discussing modern events)
4. Focus strictly on the lesson topic: {full_lesson_name}
5. Maintain the 30-section structure for the 30-slide presentation
6. Write in continuous paragraphs (NOT bullet points)
7. Each section should be 3-5 sentences
8. Ensure all facts are directly supported by the evidence provided

Generate the revised teacher's guide now:
""".strip()
            
            # Regenerate content with corrections
            summary = await generate(revision_prompt, max_tokens=2000)
            
            # Re-check the revised content
            print(f"[Planner] Re-checking revised content (attempt {revision_attempt})...")
            fact_check_result = await fact_checker_agent(summary, evidence)
            print(fact_check_result)
        
        # Final verdict after revision loop
        if "GO/NO-GO: GO" in fact_check_result:
            if revision_attempt > 0:
                print(f"[Planner] ✅ Content verified after {revision_attempt} revision(s)!")
                shared_context["fact_check_stats"].append({
                    "lesson": full_lesson_name,
                    "verdict": f"GO (revised {revision_attempt}x)",
                    "details": fact_check_result
                })
            else:
                print(f"[Planner] ✅ Content verified on first attempt")
                shared_context["fact_check_stats"].append({
                    "lesson": full_lesson_name,
                    "verdict": "GO",
                    "details": fact_check_result
                })
        else:
            print(f"[Planner] ⚠️ Content still has issues after {revision_attempt} revision attempts, using best available version")
            shared_context["fact_check_stats"].append({
                "lesson": full_lesson_name,
                "verdict": f"WARNING (tried {revision_attempt} revisions)",
                "details": fact_check_result
            })
    
    # Add only relevant sources to shared context
    for source in temp_sources:
        if source["title"] not in irrelevant_sources:
            shared_context["sources"].append(source)
    
    # Add to shared context (keep last 2 lessons to avoid token overflow)
    shared_context["lesson_summaries"].append(f"{full_lesson_name}:\n{summary[:1000]}...")
    if len(shared_context["lesson_summaries"]) > 2:
        shared_context["lesson_summaries"].pop(0)
    
    # Store full summary for quiz generation
    shared_context["full_summaries"].append(summary)
    
    # 3. Create DOCX for Summary
    docx_filename = f"{_slugify_title(full_lesson_name)}.docx"
    docx_created = create_docx(docx_filename, full_lesson_name, summary)
    docx_path = OUTPUT_DIR / docx_filename if docx_created else None

    # 4. Generate Slide Structure (Keywords only) - with context from previous slides
    print(f"[Planner] Designing slides for {full_lesson_name}...")
    slides_json_str = await generate(
        _slide_generation_prompt(
            full_lesson_name, 
            summary, 
            SLIDE_TARGET,
            shared_context["slide_titles"]
        ),
        max_tokens=4000,  
        temperature=0.3
    )
    slides_data = _safe_json_loads(slides_json_str)
    slides_list = slides_data.get("slides", [])

    if not slides_list:
        print(f"[Planner] ERROR: Failed to generate slides. Raw output: {slides_json_str[:500]}...")
        slides_list = [{"title": "Error generating slides", "bullets": ["Check logs."], "notes": "Slide generation failed."}]
    else:
        # Extract notes from Teacher's Guide and add to slides
        slides_list = _extract_notes_from_summary(summary, slides_list)
    
    # Track slide titles
    for slide in slides_list:
        shared_context["slide_titles"].append(slide.get("title", ""))

    # 5. Dispatch to PPT Agent
    ppt_filename = build_ppt_filename(full_lesson_name)
    ppt_task_id = str(uuid.uuid4())
    
    await ppt_queue.put({
        "id": ppt_task_id,
        "to": "ppt",
        "payload": {
            "title": full_lesson_name,
            "slides": slides_list, 
            "filename": ppt_filename
        }
    })

    ppt_path = OUTPUT_DIR / ppt_filename
    ppt_ready = False
    for _ in range(60): 
        if ppt_path.exists():
            ppt_ready = True
            break
        await asyncio.sleep(0.5)
    
    return {
        "lesson_name": full_lesson_name,
        "topics": topics,
        "ppt_path": str(ppt_path) if ppt_ready else None,
        "docx_path": str(docx_path) if docx_path else None
    }


# -------------------------------------------------------------------------
#  MAIN PLANNER LOGIC
# -------------------------------------------------------------------------

async def _orchestrate_history_package(request: str):
    
    # 1. Parsing
    intent_prompt = f"""
Extract topic, number of lessons, and student age from: "{request}".
Return JSON: {{"topic": "...", "num_lessons": 3, "age": 16}}
Default num_lessons to 1 if not specified.
Default age to 16 if not specified.
Age should be between 14 and 18.
""".strip()
    
    intent_raw = await generate(intent_prompt)
    intent = _safe_json_loads(intent_raw)
    topic = intent.get("topic", request)
    try:
        num_lessons = int(intent.get("num_lessons", 1))
    except:
        num_lessons = 1
    try:
        age = int(intent.get("age", 16))
        age = max(14, min(18, age))  # Clamp between 14-18
    except:
        age = 16
        
    print(f"[Planner] Plan: {num_lessons} lessons on '{topic}'")

    # 2. Planning
    plan_raw = await generate(_plan_lessons_prompt(topic, num_lessons))
    unit_plan = _safe_json_loads(plan_raw)
    
    unit_title = unit_plan.get("unit_title", topic)
    lessons = unit_plan.get("lessons", [])

    if not lessons:
        return "❌ Failed to generate a valid lesson plan."

    # 3. Build a readable plan summary
    plan_summary = f"**Unit:** {unit_title}\n**Lessons:**\n"
    for lesson in lessons:
        l_num = lesson.get("lesson_number", "?")
        l_title = lesson.get("title", "Untitled")
        plan_summary += f"  {l_num}. {l_title}\n"

    # 4. Initialize shared context to track what's been covered
    shared_context = {
        "evidence_cache": {},  # Cache Wikipedia results
        "lesson_summaries": [],  # Track previous lesson content (truncated)
        "slide_titles": [],  # Track all slide titles to avoid duplicates
        "full_summaries": [],  # Full lesson summaries for quiz generation
        "sources": [],  # Track all sources used
        "fact_check_stats": []  # Track fact checking statistics
    }

    # 5. Execute Lessons with shared context
    lesson_results = []
    for lesson in lessons:
        lesson_data = await _process_lesson(lesson, unit_title, shared_context)
        lesson_results.append(lesson_data)

    # 6. Generate Quiz
    print(f"[Planner] Generating quiz for {age}-year-olds: {unit_title}")
    quiz_data = await generate_quiz(unit_title, shared_context["full_summaries"], age)
    quiz_text = format_quiz_for_docx(unit_title, quiz_data)
    
    # Save quiz to DOCX
    quiz_filename = f"quiz_{_slugify_title(unit_title)}.docx"
    quiz_path = None
    if docx:
        try:
            doc = docx.Document()
            doc.add_heading(f"Quiz: {unit_title}", 0)
            doc.add_paragraph(f"Age Group: {age} years old")
            doc.add_paragraph("")
            
            # Add questions
            for i, q in enumerate(quiz_data.get("questions", []), 1):
                doc.add_paragraph(f"{i}. {q}")
            
            quiz_path = OUTPUT_DIR / quiz_filename
            doc.save(quiz_path)
            print(f"[Planner] Quiz saved to: {quiz_path}")
        except Exception as e:
            print(f"[Planner] Error saving quiz: {e}")

    # 7. Format Discord Output
    output_lines = [
        f"**📝 Request:** {request}",
        "",
        f"**📋 Plan made by planner agent:**",
        plan_summary,
        "",
        "**📦 Here are your files:**"
    ]

    for result in lesson_results:
        lesson_name = result["lesson_name"]
        output_lines.append(f"- {lesson_name}")
        
        if result["ppt_path"]:
            output_lines.append(f"  __FILE__:{result['ppt_path']}")
        if result["docx_path"]:
            output_lines.append(f"  __FILE__:{result['docx_path']}")
    
    # Add quiz file
    if quiz_path:
        output_lines.append(f"- Quiz")
        output_lines.append(f"  __FILE__:{str(quiz_path)}")
    
    # Generate sources document
    if shared_context["sources"]:
        sources_filename = f"sources_{_slugify_title(unit_title)}.docx"
        if create_sources_document(sources_filename, unit_title, shared_context["sources"]):
            sources_path = OUTPUT_DIR / sources_filename
            output_lines.append(f"- Research Sources Document")
            output_lines.append(f"  __FILE__:{str(sources_path)}")
    
    # Add fact check summary if enabled
    if FACT_CHECK_ENABLED and shared_context["fact_check_stats"]:
        output_lines.append("")
        output_lines.append("**🛡️ Fact Checking Summary:**")
        
        # Count different verdict types
        go_first_attempt = sum(1 for s in shared_context["fact_check_stats"] if s["verdict"] == "GO")
        go_revised = sum(1 for s in shared_context["fact_check_stats"] if "revised" in s["verdict"].lower() and "GO" in s["verdict"])
        warnings = sum(1 for s in shared_context["fact_check_stats"] if "WARNING" in s["verdict"])
        
        # Display summary
        if go_first_attempt > 0:
            output_lines.append(f"✅ {go_first_attempt} lesson(s) verified on first attempt")
        if go_revised > 0:
            output_lines.append(f"✅ {go_revised} lesson(s) revised and approved")
        if warnings > 0:
            output_lines.append(f"⚠️ {warnings} lesson(s) have warnings")
        
        # Show details for each lesson
        output_lines.append("")
        output_lines.append("**Details by lesson:**")
        for stat in shared_context["fact_check_stats"]:
            lesson_name = stat["lesson"].replace("Lesson ", "L")  # Shorten for readability
            verdict = stat["verdict"]
            
            if verdict == "GO":
                output_lines.append(f"  • {lesson_name}: ✅ Approved")
            elif "revised" in verdict.lower() and "GO" in verdict:
                # Extract revision count if present
                import re
                rev_match = re.search(r'revised (\d+)x', verdict)
                rev_count = rev_match.group(1) if rev_match else "1"
                output_lines.append(f"  • {lesson_name}: ✅ Revised {rev_count}x and approved")
            else:
                output_lines.append(f"  • {lesson_name}: ⚠️ {verdict}")

    return "\n".join(output_lines)


async def planner_agent():
    print("[Planner] Enhanced History Planner Started.")
    
    while True:
        task_str = await task_queue.get()
        print(f"\n[Planner] Task received: {task_str}")

        try:
            output = await _orchestrate_history_package(task_str)
        except Exception as e:
            output = f"❌ Error during planning execution: {e}"
            print(output) # log to console

        await result_queue.put(output)
        task_queue.task_done()
